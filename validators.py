import re
import logging
from typing import List, Dict, Tuple
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class ResumeValidator:
    """Soft-gated validation that detects violations without blocking inference."""
    
    ACCEPTABLE_FONTS = {'sans-serif', 'calibri', 'arial', 'helvetica', 'verdana', 'tahoma'}
    ENTRY_LEVEL_PAGE_TARGET = 1
    
    def __init__(self):
        self.violations = []
    
    def validate_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text and metadata from PDF, return violations."""
        violations = []
        text = ""
        page_count = 0
        fonts = set()
        links = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
                    
                    # Extract fonts
                    chars = page.chars
                    for char in chars:
                        if char.get('fontname'):
                            font_name = char['fontname'].lower()
                            fonts.add(font_name)
                    
                    # Extract links
                    links_on_page = getattr(page, 'hyperlinks', [])
                    for link in links_on_page:
                        if link.get('uri'):
                            links.append(link['uri'])
        
        except Exception as e:
            violations.append(f"PDF extraction error: {str(e)}")
            logger.error(f"PDF extraction failed: {e}")
        
        # Check if text was extracted
        if not text or len(text.strip()) < 10:
            violations.append("Warning: Very little or no text extracted from PDF. File may be image-based or corrupted.")
        
        # Check page count
        if page_count > self.ENTRY_LEVEL_PAGE_TARGET:
            violations.append(f"Page count violation: {page_count} pages (target: {self.ENTRY_LEVEL_PAGE_TARGET} for entry-level)")
        
        # Check fonts
        invalid_fonts = []
        for font in fonts:
            font_lower = font.lower()
            is_valid = any(acceptable in font_lower for acceptable in self.ACCEPTABLE_FONTS)
            if not is_valid:
                invalid_fonts.append(font)
        
        if invalid_fonts:
            violations.append(f"Font violation: Found non-Sans-Serif/Calibri fonts: {', '.join(set(invalid_fonts))}")
        
        # Check links
        invalid_links = []
        for link in links:
            if not link.startswith(('mailto:', 'https://', 'http://')):
                invalid_links.append(link)
        
        if invalid_links:
            violations.append(f"Link violation: Links must use mailto: or https:// format. Found: {', '.join(invalid_links[:5])}")
        
        metadata = {
            'page_count': page_count,
            'fonts': list(fonts),
            'links': links,
            'violations': violations
        }
        
        return text.strip(), metadata
    
    def validate_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text and metadata from DOCX, return violations."""
        violations = []
        text = ""
        fonts = set()
        links = []
        
        try:
            doc = Document(file_path)
            
            # Extract text and fonts
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
                for run in paragraph.runs:
                    if run.font.name:
                        fonts.add(run.font.name.lower())
            
            # Extract links (hyperlinks in DOCX)
            for paragraph in doc.paragraphs:
                for hyperlink in paragraph.hyperlinks:
                    if hyperlink.address:
                        links.append(hyperlink.address)
            
            # Count pages (approximate: ~500 words per page)
            word_count = len(text.split())
            page_count = max(1, (word_count // 500) + 1)
            
        except Exception as e:
            violations.append(f"DOCX extraction error: {str(e)}")
            logger.error(f"DOCX extraction failed: {e}")
            page_count = 0
        
        # Check if text was extracted
        if not text or len(text.strip()) < 10:
            violations.append("Warning: Very little or no text extracted from DOCX. File may be corrupted.")
        
        # Check page count
        if page_count > self.ENTRY_LEVEL_PAGE_TARGET:
            violations.append(f"Page count violation: Estimated {page_count} pages (target: {self.ENTRY_LEVEL_PAGE_TARGET} for entry-level)")
        
        # Check fonts
        invalid_fonts = []
        for font in fonts:
            font_lower = font.lower()
            is_valid = any(acceptable in font_lower for acceptable in self.ACCEPTABLE_FONTS)
            if not is_valid:
                invalid_fonts.append(font)
        
        if invalid_fonts:
            violations.append(f"Font violation: Found non-Sans-Serif/Calibri fonts: {', '.join(set(invalid_fonts))}")
        
        # Check links
        invalid_links = []
        for link in links:
            if not link.startswith(('mailto:', 'https://', 'http://')):
                invalid_links.append(link)
        
        if invalid_links:
            violations.append(f"Link violation: Links must use mailto: or https:// format. Found: {', '.join(invalid_links[:5])}")
        
        metadata = {
            'page_count': page_count,
            'fonts': list(fonts),
            'links': links,
            'violations': violations
        }
        
        return text.strip(), metadata
    
    def format_violations_for_prompt(self, violations: List[str]) -> str:
        """Format violations to inject into model context."""
        if not violations:
            return "No formatting violations detected."
        
        violation_text = "DETECTED FORMATTING VIOLATIONS:\n"
        for i, violation in enumerate(violations, 1):
            violation_text += f"{i}. {violation}\n"
        
        violation_text += "\nThese violations MUST impact the overall_score, ats_friendly status, weaknesses, and fixes sections."
        
        return violation_text

